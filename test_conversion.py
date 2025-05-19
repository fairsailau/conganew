"""
Test script for conversion and export functionality
"""
import os
import sys
import tempfile
import docx
from app.parser import CongaTemplateParser
from app.converter import ConversionEngine
from app.validator import ValidationEngine
from app.exporter import DocxExporter


def test_conversion_with_sample(sample_path, output_dir=None):
    """
    Test conversion and export with a sample template
    
    Args:
        sample_path: Path to the sample Conga template
        output_dir: Directory to save the output (optional)
    
    Returns:
        Path to the converted file
    """
    print(f"Testing conversion with sample: {sample_path}")
    
    # Create output directory if not provided
    if not output_dir:
        output_dir = tempfile.mkdtemp()
        print(f"Created temporary output directory: {output_dir}")
    else:
        os.makedirs(output_dir, exist_ok=True)
        print(f"Using output directory: {output_dir}")
    
    # Parse the template
    parser = CongaTemplateParser(docx_file_path=sample_path)
    conga_tags = parser.parse()
    doc = parser.get_document()
    
    print(f"Found {len(conga_tags)} Conga tags in the template")
    
    # Convert the template
    converter = ConversionEngine()
    converted_doc = converter.convert_template(doc, conga_tags)
    
    # Export the converted document
    output_filename = f"converted_{os.path.basename(sample_path)}"
    output_path = os.path.join(output_dir, output_filename)
    
    exporter = DocxExporter()
    exported_path = exporter.export_docx(converted_doc, output_path)
    
    print(f"Exported converted template to: {exported_path}")
    
    # Validate the conversion
    validator = ValidationEngine()
    
    # Extract text content for validation
    original_content = "\n".join([p.text for p in doc.paragraphs])
    converted_content = "\n".join([p.text for p in converted_doc.paragraphs])
    
    validation_results = validator.validate_conversion(
        original_content, 
        converted_content,
        conga_tags
    )
    
    print("Validation results:")
    print(f"- Syntax valid: {validation_results['syntax_valid']}")
    print(f"- Completeness: {validation_results['completeness'] * 100:.1f}%")
    
    if validation_results['errors']:
        print(f"- Errors: {len(validation_results['errors'])}")
        for error in validation_results['errors']:
            print(f"  - {error['type']}")
    
    if validation_results['warnings']:
        print(f"- Warnings: {len(validation_results['warnings'])}")
        for warning in validation_results['warnings']:
            print(f"  - {warning['type']}")
    
    # Verify tag replacement
    print("\nVerifying tag replacement and formatting preservation...")
    
    # Load the exported document to verify
    exported_doc = docx.Document(exported_path)
    
    # Check for Conga tags that should have been replaced
    conga_patterns = ['&=', '{IF', '{TABLE', '{END']
    found_conga_tags = []
    
    for paragraph in exported_doc.paragraphs:
        for pattern in conga_patterns:
            if pattern in paragraph.text:
                found_conga_tags.append(f"{pattern} in: {paragraph.text}")
    
    if found_conga_tags:
        print("WARNING: Found potential unconverted Conga tags:")
        for tag in found_conga_tags:
            print(f"- {tag}")
    else:
        print("No unconverted Conga tags found in the document")
    
    # Check for Box DocGen tags
    box_patterns = ['{{', '}}', '{{#', '{{/']
    found_box_tags = []
    
    for paragraph in exported_doc.paragraphs:
        for pattern in box_patterns:
            if pattern in paragraph.text:
                found_box_tags.append(paragraph.text)
                break
    
    if found_box_tags:
        print(f"\nFound {len(found_box_tags)} paragraphs with Box DocGen tags:")
        for i, tag in enumerate(found_box_tags[:5]):  # Show first 5 examples
            print(f"- {tag}")
        
        if len(found_box_tags) > 5:
            print(f"  ... and {len(found_box_tags) - 5} more")
    else:
        print("No Box DocGen tags found in the document")
    
    return exported_path


if __name__ == "__main__":
    if len(sys.argv) < 2:
        print("Usage: python test_conversion.py <path_to_sample_docx> [output_directory]")
        sys.exit(1)
    
    sample_path = sys.argv[1]
    output_dir = sys.argv[2] if len(sys.argv) > 2 else None
    
    test_conversion_with_sample(sample_path, output_dir)
