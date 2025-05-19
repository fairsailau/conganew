"""
Parser module for extracting Conga tags from DOCX files
"""
import re
import docx
from typing import Dict, List, Tuple, Any, Optional


class CongaTemplateParser:
    """
    Parser for Conga template DOCX files
    Extracts and classifies Conga tags for conversion
    """
    
    def __init__(self, docx_file_path: str = None, docx_file_obj = None):
        """
        Initialize the parser with either a file path or file object
        
        Args:
            docx_file_path: Path to the DOCX file
            docx_file_obj: File object (used when uploading via Streamlit)
        """
        self.docx_file_path = docx_file_path
        self.docx_file_obj = docx_file_obj
        self.doc = None
        self.text_content = ""
        self.tags = []
        self.tag_locations = []  # Store paragraph/run locations for replacement
        
    def parse(self) -> List[Dict[str, Any]]:
        """
        Parse the DOCX file and extract Conga tags
        
        Returns:
            List of dictionaries containing tag information
        """
        if self.docx_file_path:
            self.doc = docx.Document(self.docx_file_path)
        else:
            self.doc = docx.Document(self.docx_file_obj)
            
        self._extract_text_and_locations()
        self._identify_tags()
        return self.tags
    
    def get_document(self) -> docx.Document:
        """
        Get the parsed document object
        
        Returns:
            docx.Document object
        """
        return self.doc
        
    def _extract_text_and_locations(self) -> None:
        """
        Extract all text from the document and track locations
        """
        full_text = []
        
        # Process paragraphs
        for i, paragraph in enumerate(self.doc.paragraphs):
            full_text.append(paragraph.text)
            
            # Track tag locations in paragraphs
            for j, run in enumerate(paragraph.runs):
                if any(pattern in run.text for pattern in ['&=', '{{', '}}', '{IF', '{TABLE', '{END']):
                    self.tag_locations.append({
                        'type': 'paragraph',
                        'paragraph_index': i,
                        'run_index': j,
                        'text': run.text,
                        'original_run': run  # Store reference to the original run
                    })
        
        # Process tables
        for table_idx, table in enumerate(self.doc.tables):
            for row_idx, row in enumerate(table.rows):
                for cell_idx, cell in enumerate(row.cells):
                    for para_idx, paragraph in enumerate(cell.paragraphs):
                        full_text.append(paragraph.text)
                        
                        # Track tag locations in tables
                        for run_idx, run in enumerate(paragraph.runs):
                            if any(pattern in run.text for pattern in ['&=', '{{', '}}', '{IF', '{TABLE', '{END']):
                                self.tag_locations.append({
                                    'type': 'table',
                                    'table_index': table_idx,
                                    'row_index': row_idx,
                                    'cell_index': cell_idx,
                                    'paragraph_index': para_idx,
                                    'run_index': run_idx,
                                    'text': run.text,
                                    'original_run': run  # Store reference to the original run
                                })
                    
        self.text_content = "\n".join(full_text)
        
    def _identify_tags(self) -> None:
        """
        Identify and classify Conga tags in the document
        """
        # Patterns for different Conga tag types
        patterns = {
            'merge_field': r'&=([A-Za-z0-9._]+)',
            'curly_brace_field': r'\{\{([^}]+)\}\}',
            'conditional': r'\{IF\s+"([^"]+)"\s+([^}]+)\}',
            'table_start': r'\{TABLE\s+([^}]+)\}',
            'table_end': r'\{END\s+([^}]+)\}'
        }
        
        for tag_type, pattern in patterns.items():
            for match in re.finditer(pattern, self.text_content):
                self.tags.append({
                    'type': tag_type,
                    'full_match': match.group(0),
                    'groups': [match.group(i) for i in range(1, match.lastindex + 1)] if match.lastindex else [],
                    'position': match.span(),
                    'location': self._find_tag_location(match.group(0))
                })
        
        # Sort tags by position in document
        self.tags.sort(key=lambda x: x['position'][0])
    
    def _find_tag_location(self, tag_text: str) -> Optional[Dict[str, Any]]:
        """
        Find the location of a tag in the document
        
        Args:
            tag_text: The tag text to find
            
        Returns:
            Dictionary with location information or None if not found
        """
        for location in self.tag_locations:
            if tag_text in location['text']:
                return location
        
        return None
