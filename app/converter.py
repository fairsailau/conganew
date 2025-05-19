"""
Conversion engine for transforming Conga templates to Box DocGen templates
"""
import re
from typing import Dict, List, Any, Optional, Tuple
import docx
from box_ai_client import BoxAIClient


class ConversionEngine:
    """
    Engine for converting Conga template tags to Box DocGen format
    """
    
    def __init__(self, box_ai_client: Optional[BoxAIClient] = None, mapping_rules: Optional[Dict] = None):
        """
        Initialize the conversion engine
        
        Args:
            box_ai_client: BoxAIClient instance for AI-assisted conversion
            mapping_rules: Custom mapping rules (optional)
        """
        self.box_ai_client = box_ai_client
        self.mapping_rules = mapping_rules or self._load_default_rules()
        
    def _load_default_rules(self) -> Dict[str, Any]:
        """
        Load default mapping rules for Conga to Box DocGen conversion
        
        Returns:
            Dictionary of mapping rules
        """
        return {
            'merge_field': {
                # Common Salesforce field mappings
                '&=Account.Name': 'account.name',
                '&=Contact.Name': 'contact.name',
                '&=Opportunity.Name': 'opportunity.name',
                '&=Date.Today': '{{date now format="dd-MM-yyyy"}}',
                # General pattern transformations
                'pattern': r'&=([A-Za-z0-9._]+)',
                'replacement': r'{{\1}}'
            },
            'curly_brace_field': {
                # Direct mapping for curly brace fields
                'pattern': r'\{\{([^}]+)\}\}',
                'replacement': r'{{\1}}',
                # Format specifiers
                'date_format': r'\\@\s*([^}]+)',
                'date_replacement': r' format="\1"'
            },
            'conditional': {
                'pattern': r'\{IF\s+"([^"]+)"\s*=\s*"([^"]+)"\s+"([^"]+)"\s+"([^"]+)"\}',
                'replacement': r'{{#eq \1 "\2"}}\3{{else}}\4{{/eq}}',
                'pattern_gt': r'\{IF\s+"([^"]+)"\s*>\s*"([^"]+)"\s+"([^"]+)"\s+"([^"]+)"\}',
                'replacement_gt': r'{{#gt \1 \2}}\3{{else}}\4{{/gt}}',
                'pattern_lt': r'\{IF\s+"([^"]+)"\s*<\s*"([^"]+)"\s+"([^"]+)"\s+"([^"]+)"\}',
                'replacement_lt': r'{{#lt \1 \2}}\3{{else}}\4{{/lt}}'
            },
            'table': {
                'pattern': r'\{TABLE\s+([^}]+)\}',
                'replacement': r'{{#each \1}}',
                'end_pattern': r'\{END\s+([^}]+)\}',
                'end_replacement': r'{{/each}}'
            }
        }
        
    def convert_template(self, doc: docx.Document, conga_tags: List[Dict[str, Any]]) -> docx.Document:
        """
        Convert Conga template to Box DocGen format
        
        Args:
            doc: docx.Document object containing the template
            conga_tags: List of Conga tags extracted from the template
            
        Returns:
            docx.Document with converted Box DocGen tags
        """
        # Process each tag by type
        for tag in conga_tags:
            if tag['location'] is None:
                continue
                
            converted_tag = self._convert_tag(tag)
            if converted_tag:
                self._replace_tag_in_document(doc, tag, converted_tag)
        
        return doc
    
    def _convert_tag(self, tag: Dict[str, Any]) -> Optional[str]:
        """
        Convert a single Conga tag to Box DocGen format
        
        Args:
            tag: Dictionary containing tag information
            
        Returns:
            Converted tag string or None if conversion failed
        """
        tag_type = tag['type']
        full_match = tag['full_match']
        
        if tag_type == 'merge_field':
            return self._convert_merge_field(full_match)
        elif tag_type == 'curly_brace_field':
            return self._convert_curly_brace_field(full_match)
        elif tag_type == 'conditional':
            return self._convert_conditional(full_match)
        elif tag_type == 'table_start':
            return self._convert_table_start(full_match)
        elif tag_type == 'table_end':
            return self._convert_table_end(full_match)
        
        # If no direct conversion is possible, try AI-assisted conversion
        if self.box_ai_client:
            return self._ai_assisted_conversion(tag)
            
        return None
    
    def _convert_merge_field(self, tag_text: str) -> str:
        """
        Convert Conga merge field to Box DocGen format
        
        Args:
            tag_text: Original Conga merge field tag
            
        Returns:
            Converted Box DocGen tag
        """
        # Check for direct mapping
        if tag_text in self.mapping_rules['merge_field']:
            return self.mapping_rules['merge_field'][tag_text]
        
        # Apply pattern replacement
        pattern = self.mapping_rules['merge_field']['pattern']
        replacement = self.mapping_rules['merge_field']['replacement']
        
        # Convert field name format (e.g., Account.Name to account.name)
        converted = re.sub(pattern, replacement, tag_text)
        if converted:
            # Convert to lowercase for Box DocGen convention
            parts = converted.split('.')
            converted = '.'.join([parts[0].lower()] + parts[1:])
            
        return converted
    
    def _convert_curly_brace_field(self, tag_text: str) -> str:
        """
        Convert Conga curly brace field to Box DocGen format
        
        Args:
            tag_text: Original Conga curly brace field tag
            
        Returns:
            Converted Box DocGen tag
        """
        # Handle date format specifiers
        date_format_pattern = self.mapping_rules['curly_brace_field']['date_format']
        date_format_match = re.search(date_format_pattern, tag_text)
        
        if date_format_match:
            # Extract the field name and format
            field_pattern = r'\{\{([^\\@]+)'
            field_match = re.search(field_pattern, tag_text)
            
            if field_match:
                field_name = field_match.group(1).strip()
                format_spec = date_format_match.group(1)
                
                # Convert to Box DocGen date format
                return f"{{{{date {field_name.lower()} format=\"{format_spec}\"}}}}"
        
        # For standard fields, convert to lowercase
        field_pattern = r'\{\{([^}]+)\}\}'
        field_match = re.search(field_pattern, tag_text)
        
        if field_match:
            field_name = field_match.group(1).strip()
            # Convert to lowercase for Box DocGen convention
            return f"{{{{{field_name.lower()}}}}}"
            
        return tag_text
    
    def _convert_conditional(self, tag_text: str) -> Optional[str]:
        """
        Convert Conga conditional to Box DocGen format
        
        Args:
            tag_text: Original Conga conditional tag
            
        Returns:
            Converted Box DocGen conditional or None if conversion failed
        """
        # Try equality condition
        eq_pattern = self.mapping_rules['conditional']['pattern']
        eq_replacement = self.mapping_rules['conditional']['replacement']
        eq_match = re.match(eq_pattern, tag_text)
        
        if eq_match:
            field = eq_match.group(1).lower()
            value = eq_match.group(2)
            true_value = eq_match.group(3)
            false_value = eq_match.group(4)
            
            return f"{{{{#eq {field} \"{value}\"}}}}{true_value}{{{{else}}}}{false_value}{{{{/eq}}}}"
        
        # Try greater than condition
        gt_pattern = self.mapping_rules['conditional']['pattern_gt']
        gt_match = re.match(gt_pattern, tag_text)
        
        if gt_match:
            field = gt_match.group(1).lower()
            value = gt_match.group(2)
            true_value = gt_match.group(3)
            false_value = gt_match.group(4)
            
            return f"{{{{#gt {field} {value}}}}}{true_value}{{{{else}}}}{false_value}{{{{/gt}}}}"
        
        # Try less than condition
        lt_pattern = self.mapping_rules['conditional']['pattern_lt']
        lt_match = re.match(lt_pattern, tag_text)
        
        if lt_match:
            field = lt_match.group(1).lower()
            value = lt_match.group(2)
            true_value = lt_match.group(3)
            false_value = lt_match.group(4)
            
            return f"{{{{#lt {field} {value}}}}}{true_value}{{{{else}}}}{false_value}{{{{/lt}}}}"
        
        # If no direct pattern match, try AI-assisted conversion
        if self.box_ai_client:
            prompt = f"""
            Convert this Conga template conditional tag to Box DocGen format:
            
            {tag_text}
            
            Return only the Box DocGen equivalent without explanation.
            """
            
            response = self.box_ai_client.generate_text(prompt)
            if 'answer' in response:
                return response['answer'].strip()
                
        return None
    
    def _convert_table_start(self, tag_text: str) -> str:
        """
        Convert Conga table start tag to Box DocGen format
        
        Args:
            tag_text: Original Conga table start tag
            
        Returns:
            Converted Box DocGen tag
        """
        pattern = self.mapping_rules['table']['pattern']
        replacement = self.mapping_rules['table']['replacement']
        
        match = re.match(pattern, tag_text)
        if match:
            collection = match.group(1).strip()
            # Extract collection name and convert to lowercase
            collection_parts = collection.split('=')
            if len(collection_parts) > 1:
                collection_name = collection_parts[1].strip().lower()
                return f"{{{{#each {collection_name}}}}}"
        
        # Default replacement
        return re.sub(pattern, replacement, tag_text)
    
    def _convert_table_end(self, tag_text: str) -> str:
        """
        Convert Conga table end tag to Box DocGen format
        
        Args:
            tag_text: Original Conga table end tag
            
        Returns:
            Converted Box DocGen tag
        """
        pattern = self.mapping_rules['table']['end_pattern']
        replacement = self.mapping_rules['table']['end_replacement']
        
        return re.sub(pattern, replacement, tag_text)
    
    def _ai_assisted_conversion(self, tag: Dict[str, Any]) -> Optional[str]:
        """
        Use Box AI for complex conversion scenarios
        
        Args:
            tag: Dictionary containing tag information
            
        Returns:
            AI-generated conversion or None if not available
        """
        if not self.box_ai_client:
            return None
            
        prompt = f"""
        Convert this Conga template tag to Box DocGen format:
        
        Conga tag: {tag['full_match']}
        Tag type: {tag['type']}
        
        Return only the Box DocGen equivalent without explanation.
        """
        
        response = self.box_ai_client.generate_text(prompt)
        if 'answer' in response:
            return response['answer'].strip()
            
        return None
    
    def _replace_tag_in_document(self, doc: docx.Document, tag: Dict[str, Any], 
                                converted_tag: str) -> None:
        """
        Replace a tag in the document with its converted version
        
        Args:
            doc: docx.Document object
            tag: Dictionary containing tag information
            converted_tag: Converted Box DocGen tag
        """
        location = tag['location']
        if not location:
            return
            
        if location['type'] == 'paragraph':
            paragraph = doc.paragraphs[location['paragraph_index']]
            run = paragraph.runs[location['run_index']]
            
            # Replace the tag while preserving formatting
            run.text = run.text.replace(tag['full_match'], converted_tag)
            
        elif location['type'] == 'table':
            table = doc.tables[location['table_index']]
            cell = table.rows[location['row_index']].cells[location['cell_index']]
            paragraph = cell.paragraphs[location['paragraph_index']]
            run = paragraph.runs[location['run_index']]
            
            # Replace the tag while preserving formatting
            run.text = run.text.replace(tag['full_match'], converted_tag)
