"""
Main Streamlit application for Conga to Box DocGen template conversion
"""
import io
import json
import os
import tempfile
import zipfile
from pathlib import Path
from typing import Dict, List, Optional, Any, Union, Tuple

import streamlit as st
from docx import Document

# Box SDK imports
from boxsdk import Client
from boxsdk.exception import BoxException

# Local application imports
# Import in the order of dependency to avoid circular imports

# First, import modules that don't have internal dependencies
try:
    from app.box_ai_client import BoxAIClient, BoxAIClientError, BoxAuthError, AuthMethod
except ImportError:
    from box_ai_client import BoxAIClient, BoxAIClientError, BoxAuthError, AuthMethod

try:
    from app.prompt_builder import PromptBuilder, ConversionContext
except ImportError:
    from prompt_builder import PromptBuilder, ConversionContext

try:
    from app.response_parser import AIResponseParser
except ImportError:
    from response_parser import AIResponseParser

# Then import modules that might depend on the above
try:
    from app.exporter import DocxExporter
except ImportError:
    from exporter import DocxExporter

try:
    from app.query_loader import CongaQueryLoader
except ImportError:
    from query_loader import CongaQueryLoader

try:
    from app.parser import CongaTemplateParser
except ImportError:
    from parser import CongaTemplateParser

# Import conversion_engine last as it might depend on other modules
try:
    from app.conversion_engine import ConversionEngine
except ImportError:
    from conversion_engine import ConversionEngine

try:
    from .schema_loader import JSONSchemaLoader
except ImportError:
    from schema_loader import JSONSchemaLoader

try:
    from .template_generator import DocGenTemplateGenerator
except ImportError:
    from template_generator import DocGenTemplateGenerator

try:
    from .validation_engine import ValidationEngine
except ImportError:
    from validation_engine import ValidationEngine

def initialize_session_state():
    """Initialize session state variables"""
    if 'schema_data' not in st.session_state:
        st.session_state.schema_data = None
    if 'conversion_results' not in st.session_state:
        st.session_state.conversion_results = None
    if 'error' not in st.session_state:
        st.session_state.error = None
    if 'success' not in st.session_state:
        st.session_state.success = None
    if 'converted_files' not in st.session_state:
        st.session_state.converted_files = []
    if 'processing' not in st.session_state:
        st.session_state.processing = False

def get_auth_config():
    """Get authentication configuration from Streamlit Secrets"""
    if 'BOX_DEVELOPER_TOKEN' not in st.secrets:
        st.error("Box Developer Token not found in Streamlit Secrets")
        st.stop()
    
    return {
        'developer_token': st.secrets['BOX_DEVELOPER_TOKEN'],
        'auth_method': AuthMethod.DEVELOPER_TOKEN
    }

def render_sidebar():
    """Render the sidebar with configuration options"""
    with st.sidebar:
        # Schema upload
        st.subheader("Schema Configuration")
        schema_file = st.file_uploader(
            "Upload Box-Salesforce JSON Schema (optional)",
            type=["json"],
            help="Upload a JSON schema file for better field mapping"
        )
        
        if schema_file:
            try:
                st.session_state.schema_data = json.load(schema_file)
                st.success("Schema loaded successfully!")
            except Exception as e:
                st.error(f"Error loading schema: {str(e)}")
        
        # Processing options
        st.subheader("Processing Options")
        use_ai = st.checkbox(
            "Enable AI-assisted conversion",
            value=True,
            help="Use Box AI for complex conversions"
        )
        
        validate_output = st.checkbox(
            "Validate output",
            value=True,
            help="Validate the converted template"
        )
        
        return use_ai, validate_output

def main():
    """Main application function."""
    st.set_page_config(
        page_title="Conga to Box DocGen Converter",
        page_icon="🔄",
        layout="wide",
        initial_sidebar_state="expanded"
    )
    
    # Initialize session state
    initialize_session_state()
    
    # Show configuration in sidebar
    use_ai, validate_output = render_sidebar()
    
    # Get authentication config
    try:
        auth_config = get_auth_config()
    except Exception as e:
        st.error(f"Authentication error: {str(e)}")
        st.stop()
    
    # Main app UI
    st.title("Conga to Box DocGen Converter")
    
    # File upload section
    st.header("Upload Conga Template")
    uploaded_file = st.file_uploader(
        "Upload your Conga template (DOCX or ZIP)",
        type=["docx", "zip"],
        accept_multiple_files=False,
        help="Upload a single DOCX file or a ZIP file containing multiple DOCX files"
    )
    
    # Process button
    if uploaded_file is not None and st.button("Convert to Box DocGen"):
        with st.spinner("Processing template..."):
            try:
                # Process the uploaded file
                process_conversion(
                    uploaded_file=uploaded_file,
                    schema_data=st.session_state.schema_data,
                    use_ai=use_ai,
                    validate_output=validate_output,
                    auth_config=auth_config
                )
                st.session_state.success = "Conversion completed successfully!"
            except Exception as e:
                st.session_state.error = f"Error during conversion: {str(e)}"
                st.error(st.session_state.error)
    
    # Show success message if conversion was successful
    if st.session_state.success:
        st.success(st.session_state.success)
    
    # Show conversion results if available
    if st.session_state.conversion_results:
        show_conversion_results()
    
    # Main content area
    st.title("Conga to Box DocGen Template Converter")
    st.markdown("""
    Upload a Conga template file or paste a Conga SOQL query to convert it to a Box DocGen template.
    """)

    # File upload or query input
    tab1, tab2 = st.tabs(["Upload Template File", "Paste SOQL Query"])

    with tab1:
        uploaded_file = st.file_uploader(
            "Upload a Conga template file (DOCX, PPTX, or XLSX)",
            type=["docx", "pptx", "xlsx"],
            key="template_uploader"
        )
        query_text = ""

    with tab2:
        query_text = st.text_area(
            "Paste your Conga SOQL query here",
            height=200,
            help="Example: SELECT Id, Name FROM Account"
        )
        uploaded_file = None

    # Custom instructions
    custom_instructions = st.text_area(
        "Custom Instructions (optional)",
        help="Provide any specific instructions for the conversion"
    )

    # Convert button
    if st.button("Convert to Box DocGen", type="primary"):
        if not uploaded_file and not query_text.strip():
            st.error("Please upload a file or paste a SOQL query")
            return

        if not st.session_state.auth_config:
            st.error("Please configure authentication first")
            return

        try:
            with st.spinner("Converting template..."):
                process_conversion(
                    uploaded_file=uploaded_file,
                    query_text=query_text,
                    schema_data=st.session_state.schema_data,
                    use_ai=use_ai,
                    validate_output=validate_output,
                    auth_config=st.session_state.auth_config,
                    custom_instructions=custom_instructions
                )
                # Refresh the page to show results
                st.rerun()
                
        except BoxAuthError as e:
            st.error(f"Authentication error: {str(e)}")
        except Exception as e:
            st.error(f"An error occurred during conversion: {str(e)}")
            st.exception(e)  # This will show the full traceback in the app


def process_conversion(
    uploaded_file: Optional[Any] = None,
    query_text: str = "",
    schema_data: Optional[Dict] = None,
    use_ai: bool = False,
    validate_output: bool = True,
    auth_config: Optional[Dict[str, Any]] = None,
    custom_instructions: str = ""
) -> None:
    """
    Process the conversion of a Conga template to Box DocGen format
    
    Args:
        uploaded_file: Uploaded template file object (can be None if using query)
        query_text: Conga SOQL query text
        schema_data: JSON schema data for field mapping
        use_ai: Whether to use Box AI for complex conversions
        validate_output: Whether to validate the converted template
        auth_config: Authentication configuration for Box API
        custom_instructions: Custom instructions for the conversion
    """
    # Initialize Box AI client if needed
    box_ai_client = None
    if use_ai:
        if not auth_config:
            st.error("Authentication configuration is required for AI-assisted conversion")
            return
            
        try:
            box_ai_client = BoxAIClient(auth_config)
            st.session_state.box_ai_client = box_ai_client
        except BoxAuthError as e:
            st.error(f"Failed to authenticate with Box: {str(e)}")
            return
        except Exception as e:
            st.error(f"Failed to initialize Box AI client: {str(e)}")
            return
    
    # Initialize other components
    schema_loader = JSONSchemaLoader(schema_data) if schema_data else None
    template_generator = DocGenTemplateGenerator()
    
    # Process template file if provided
    template_content = ""
    if uploaded_file:
        try:
            # Handle both file path and file-like objects
            if hasattr(uploaded_file, 'read'):
                doc = Document(io.BytesIO(uploaded_file.getvalue()))
            else:
                doc = Document(uploaded_file)
                
            template_content = "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
            
            # Store the uploaded file info in session state
            st.session_state.uploaded_file_name = uploaded_file.name
            
        except Exception as e:
            st.error(f"Error reading template file: {str(e)}")
            st.exception(e)
            return
    
    # Process query if provided
    query_components = {}
    if query_text and query_text.strip():
        try:
            query_loader = CongaQueryLoader(query_text)
            query_components = query_loader.get_query_components()
        except Exception as e:
            st.error(f"Error processing query: {str(e)}")
            st.exception(e)
            return
    
    # Validate that we have either a template or a query
    if not template_content and not query_components:
        st.error("Please provide either a template file or a SOQL query")
        return
    
    # Build the conversion context
    context = ConversionContext(
        template_text=template_content,
        query_text=query_text,
        schema_data=schema_data,
        custom_instructions=custom_instructions
    )
    
    # Process the conversion
    try:
        if use_ai and box_ai_client:
            # Build the prompt
            prompt_builder = PromptBuilder(context)
            prompt = prompt_builder.build_conversion_prompt()
            
            # Use AI for conversion
            response = box_ai_client.generate_text(
                prompt=prompt['user_prompt'],
                system_prompt=prompt['system_prompt'],
                max_tokens=4000
            )
            
            # Parse the AI response
            parsed_response = AIResponseParser.parse_conversion_result(response)
            converted_content = parsed_response['content']
        else:
            # Fall back to rule-based conversion
            converter = ConversionEngine(box_ai_client=box_ai_client)
            if template_content:
                converted_content = converter.convert_text(template_content)
            else:
                converted_content = ""
    
        # Generate the output document
        output_doc = template_generator.create_from_ai_output(converted_content)
        
        # Save the output to a temporary file
        with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as tmp_file:
            output_path = tmp_file.name
        
        template_generator.save_to_file(output_doc, output_path)
        
        # Validate the output if requested
        validation_results = {}
        if validate_output:
            try:
                validator = ValidationEngine(box_ai_client=box_ai_client)
                validation_results = validator.validate_conversion(
                    template_content,
                    converted_content,
                    {}
                )
            except Exception as e:
                st.warning(f"Validation failed: {str(e)}")
        
        # Store results in session state
        st.session_state.converted_doc = output_path
        st.session_state.validation_results = validation_results
        
    except Exception as e:
        st.error(f"Error during conversion: {str(e)}")
        st.exception(e)
        raise


def preview_docx(uploaded_file) -> None:
    """
    Preview a DOCX file in Streamlit
    
    Args:
        uploaded_file: Uploaded file object or file path
    """
    try:
        if hasattr(uploaded_file, 'read'):
            # Handle file upload object
            doc = docx.Document(io.BytesIO(uploaded_file.getvalue()))
        else:
            # Handle file path
            doc = docx.Document(uploaded_file)
        
        # Display paragraphs
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                st.write(paragraph.text)
        
        # Display tables
        for i, table in enumerate(doc.tables):
            st.write(f"**Table {i+1}:**")
            
            # Create a markdown table
            table_data = []
            for row in table.rows:
                row_data = [cell.text for cell in row.cells]
                table_data.append(row_data)
            
            if table_data:
                # Create markdown table
                md_table = "| " + " | ".join(table_data[0]) + " |\n"
                md_table += "|" + "|".join(["---"] * len(table_data[0])) + "|\n"
                
                for row in table_data[1:]:
                    md_table += "| " + " | ".join(row) + " |\n"
                
                st.markdown(md_table)
    
    except Exception as e:
        st.error(f"Error previewing document: {str(e)}")


def preview_docx_from_path(file_path: str) -> None:
    """
    Preview a DOCX file from a file path
    
    Args:
        file_path: Path to the DOCX file
    """
    try:
        preview_docx(file_path)
    except Exception as e:
        st.error(f"Error loading document: {str(e)}")


if __name__ == "__main__":
    main()

def show_documentation():
    """Display documentation and examples for the template conversion."""
    st.write("## 📚 Documentation & Examples")
    st.write("""
    This tool helps you convert Conga templates to Box DocGen format. Below are some examples
    of common patterns and their Box DocGen equivalents.
    """)
    
    # Create tabs for different documentation sections
    tab1, tab2, tab3, tab4 = st.tabs([
        "🔗 Merge Fields", 
        "⚡ Conditionals", 
        "📊 Tables", 
        "🔑 Authentication"
    ])
    
    # Merge Fields Tab
    with tab1:
        st.markdown("### Merge Field Mappings")
        st.code(
            """# Merge Field Mappings
'&=Account.Name' -> 'account.name'
'&=Contact.Name' -> 'contact.name'
'&=Opportunity.Name' -> 'opportunity.name'
'&=Date.Today' -> '{{date now format="dd-MM-yyyy"}}'""",
            language="markdown"
        )
    
    # Conditionals Tab
    with tab2:
        st.markdown("### Conditional Logic Mappings")
        st.code(
            """# Equality Condition
{IF "{{field}}" = "value" "true_result" "false_result"}
->
{{#eq field "value"}}true_result{{else}}false_result{{/eq}}

# Greater Than Condition
{IF "{{field}}" > "value" "true_result" "false_result"}
->
{{#gt field value}}true_result{{else}}false_result{{/gt}}""",
            language="markdown"
        )
    
    # Tables Tab
    with tab3:
        st.markdown("### Table Mappings")
        st.code(
            """# Table Start
{TABLE Group=collection_name}
->
{{#each collection_name}}

# Table End
{END collection_name}
->
{{/each}}""",
            language="markdown"
        )
    
    # Authentication Tab
    with tab4:
        st.markdown("### Authentication Methods")
        st.write(
            """#### Supported Authentication Methods:
               - Best for server-to-server communication
            
            2. **OAuth 2.0 with Client Credentials Grant**
               - For machine-to-machine authentication
               - No user interaction required
            
            3. **OAuth 2.0 with Authorization Code Grant**
               - For user-facing applications
               - Requires user login and consent
            """)


def show_validation_results(validation_results: Dict) -> None:
    """Display validation results in a user-friendly format.
    
    Args:
        validation_results: Dictionary containing validation results
    """
    if not validation_results:
        return
    
    st.subheader("🔍 Validation Results")
    
    # Display overall status
    col1, col2 = st.columns(2)
    with col1:
        status = "✅ Valid" if validation_results.get("syntax_valid", False) else "❌ Invalid"
        st.metric("Syntax Validation", status)
    
    with col2:
        completeness = validation_results.get("completeness", 0) * 100
        st.metric("Template Completeness", f"{completeness:.1f}%")
    
    # Display errors if any
    if validation_results.get("errors"):
        st.error(f"Found {len(validation_results['errors'])} errors:")
        for error in validation_results["errors"]:
            st.write(f"- {error.get('message', 'Unknown error')}")
    
    # Display warnings if any
    if validation_results.get("warnings"):
        st.warning(f"Found {len(validation_results['warnings'])} warnings:")
        for warning in validation_results["warnings"]:
            st.write(f"- {warning.get('message', 'Warning')}")
    
    # Show AI analysis if available
    if validation_results.get("ai_analysis"):
        with st.expander("🤖 AI Analysis"):
            st.write(validation_results["ai_analysis"])
    
def show_conversion_results() -> None:
    """Display the conversion results and download options."""
    st.header("🎯 Conversion Results")
    
    if not hasattr(st.session_state, 'converted_doc') or not st.session_state.converted_doc:
        st.info("No conversion results to display. Please convert a template first.")
        return
    
    # Show validation results if available
    if hasattr(st.session_state, 'validation_results') and st.session_state.validation_results:
        show_validation_results(st.session_state.validation_results)
    
    # Show download button for the converted file
    output_path = st.session_state.converted_doc
    filename = os.path.basename(output_path)
    
    with open(output_path, "rb") as file:
        st.download_button(
            label="⬇️ Download Converted Template",
            data=file,
            file_name=f"converted_{filename}",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            key="download_converted"
        )
    
    # Show preview of the converted document
    st.subheader("🔍 Preview")
    preview_docx_from_path(output_path)


def main() -> None:
    """Main application function."""
    st.set_page_config(
        page_title="Conga to Box DocGen Converter",
        page_icon="📄",
        layout="wide",
        initial_sidebar_state="expanded"
    )
    
    # Initialize session state
    if 'converted_doc' not in st.session_state:
        st.session_state.converted_doc = None
    if 'validation_results' not in st.session_state:
        st.session_state.validation_results = {}
    if 'box_ai_client' not in st.session_state:
        st.session_state.box_ai_client = None
    
    # Sidebar for authentication and settings
    with st.sidebar:
        st.title("⚙️ Settings")
        
        # Authentication method selection
        auth_method = st.selectbox(
            "Authentication Method",
            ["JWT", "OAuth 2.0 (Client Credentials)", "OAuth 2.0 (Authorization Code)"],
            index=0
        )
        
        # Authentication configuration based on selected method
        auth_config = {}
        if auth_method == "JWT":
            st.subheader("🔑 JWT Configuration")
            auth_config = {
                "auth_method": "jwt",
                "client_id": st.text_input("Client ID", ""),
                "client_secret": st.text_input("Client Secret", "", type="password"),
                "enterprise_id": st.text_input("Enterprise ID", ""),
                "jwt_key_id": st.text_input("JWT Key ID", ""),
                "rsa_private_key_data": st.text_area(
                    "RSA Private Key", 
                    placeholder="-----BEGIN ENCRYPTED PRIVATE KEY-----\n...\n-----END ENCRYPTED PRIVATE KEY-----",
                    height=200
                ),
                "rsa_private_key_passphrase": st.text_input("Passphrase (if any)", "", type="password")
            }
        elif "Client Credentials" in auth_method:
            st.subheader("🔑 OAuth 2.0 Configuration")
            auth_config = {
                "auth_method": "ccg",
                "client_id": st.text_input("Client ID", ""),
                "client_secret": st.text_input("Client Secret", "", type="password"),
                "enterprise_id": st.text_input("Enterprise ID (optional)", "")
            }
        else:
            st.subheader("🔑 OAuth 2.0 Configuration")
            auth_config = {
                "auth_method": "acg",
                "client_id": st.text_input("Client ID", ""),
                "client_secret": st.text_input("Client Secret", "", type="password"),
                "redirect_uri": st.text_input("Redirect URI", "http://localhost:8501/callback")
            }
        
        # Conversion options
        st.subheader("🛠️ Conversion Options")
        use_ai = st.checkbox("Use Box AI for complex conversions", value=True)
        validate_output = st.checkbox("Validate output template", value=True)
        
        # Documentation link
        st.markdown("---")
        st.markdown("### 📚 Documentation")
        show_documentation()
    
    # Main content area
    st.title("📄 Conga to Box DocGen Converter")
    
    # File upload section
    st.header("📤 Upload Template")
    template_file = st.file_uploader(
        "Upload a Conga template (DOCX)", 
        type=["docx"],
        help="Upload a Word document containing Conga merge fields"
    )
    
    # Query input section
    st.header("🔍 SOQL Query (Optional)")
    query_text = st.text_area(
        "Enter your SOQL query",
        height=150,
        help="Enter a SOQL query to extract data for the template"
    )
    
    # Schema input section
    st.header("📋 Schema (Optional)")
    schema_data = None
    schema_file = st.file_uploader(
        "Upload a JSON schema for field mapping (optional)",
        type=["json"]
    )
    
    if schema_file:
        try:
            schema_data = json.load(schema_file)
            st.success("Schema loaded successfully!")
        except Exception as e:
            st.error(f"Error loading schema: {str(e)}")
    
    # Custom instructions
    custom_instructions = st.text_area(
        "Custom Instructions (Optional)",
        height=100,
        help="Provide any specific instructions for the conversion"
    )
    
    # Convert button
    convert_btn = st.button("🔄 Convert Template", type="primary")
    
    # Process conversion when button is clicked
    if convert_btn:
        if template_file or query_text:
            with st.spinner("Converting template..."):
                try:
                    process_conversion(
                        template_file=template_file,
                        query_text=query_text,
                        schema_data=schema_data,
                        use_ai=use_ai,
                        validate_output=validate_output,
                        auth_config=auth_config,
                        custom_instructions=custom_instructions
                    )
                    st.success("Conversion completed successfully!")
                except Exception as e:
                    st.error(f"Error during conversion: {str(e)}")
                    st.exception(e)
        else:
            st.warning("Please upload a template file or enter a query to convert.")
    
    # Show conversion results if available
    if hasattr(st.session_state, 'converted_doc') and st.session_state.converted_doc:
        show_conversion_results()


def process_single_conversion(uploaded_file, use_ai: bool, box_token: str) -> None:
    """
    Process a single file conversion
    
    Args:
        uploaded_file: Uploaded file object
        use_ai: Whether to use Box AI for complex conversions
        box_token: Box API token
    """
    # Create a temporary file
    with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as tmp_file:
        tmp_file.write(uploaded_file.getvalue())
        tmp_path = tmp_file.name
    
    try:
        # Initialize components
        parser = CongaTemplateParser(docx_file_path=tmp_path)
        
        # Initialize Box AI client if token is provided and AI is enabled
        box_ai_client = None
        if box_token and use_ai:
            box_ai_client = BoxAIClient(box_token)
        
        # Parse the template
        conga_tags = parser.parse()
        doc = parser.get_document()
        
        # Convert the template
        converter = ConversionEngine(box_ai_client=box_ai_client)
        converted_doc = converter.convert_template(doc, conga_tags)
        
        # Export the converted document
        output_dir = tempfile.mkdtemp()
        output_path = os.path.join(output_dir, f"converted_{uploaded_file.name}")
        
        exporter = DocxExporter()
        exported_path = exporter.export_docx(converted_doc, output_path)
        
        # Validate the conversion
        validator = ValidationEngine(box_ai_client=box_ai_client)
        
        # Extract text content for validation
        original_content = "\n".join([p.text for p in doc.paragraphs])
        converted_content = "\n".join([p.text for p in converted_doc.paragraphs])
        
        validation_results = validator.validate_conversion(
            original_content, 
            converted_content,
            conga_tags
        )
        
        # Store results in session state
        st.session_state.converted_docs[uploaded_file.name] = {
            'doc': converted_doc,
            'path': exported_path
        }
        
        st.session_state.validation_results[uploaded_file.name] = validation_results
        
        # Show success message
        st.success(f"Successfully converted {uploaded_file.name}")
        
        # Preview the converted document
        with st.expander("Preview Converted Template"):
            preview_docx_from_path(exported_path)
            
    except Exception as e:
        st.error(f"Error converting {uploaded_file.name}: {str(e)}")
    finally:
        # Clean up temporary file
        if os.path.exists(tmp_path):
            os.unlink(tmp_path)


def process_batch_conversion(uploaded_files, use_ai: bool, box_token: str) -> None:
    """
    Process batch conversion of multiple files
    
    Args:
        uploaded_files: List of uploaded file objects
        use_ai: Whether to use Box AI for complex conversions
        box_token: Box API token
    """
    progress_bar = st.progress(0)
    status_text = st.empty()
    
    # Initialize Box AI client if token is provided and AI is enabled
    box_ai_client = None
    if box_token and use_ai:
        box_ai_client = BoxAIClient(box_token)
    
    # Create output directory
    output_dir = tempfile.mkdtemp()
    
    for i, uploaded_file in enumerate(uploaded_files):
        progress = (i + 1) / len(uploaded_files)
        progress_bar.progress(progress)
        status_text.text(f"Processing {i+1}/{len(uploaded_files)}: {uploaded_file.name}")
        
        try:
            # Create a temporary file
            with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as tmp_file:
                tmp_file.write(uploaded_file.getvalue())
                tmp_path = tmp_file.name
            
            # Initialize components
            parser = CongaTemplateParser(docx_file_path=tmp_path)
            
            # Parse the template
            conga_tags = parser.parse()
            doc = parser.get_document()
            
            # Convert the template
            converter = ConversionEngine(box_ai_client=box_ai_client)
            converted_doc = converter.convert_template(doc, conga_tags)
            
            # Export the converted document
            output_path = os.path.join(output_dir, f"converted_{uploaded_file.name}")
            
            exporter = DocxExporter()
            exported_path = exporter.export_docx(converted_doc, output_path)
            
            # Validate the conversion
            validator = ValidationEngine(box_ai_client=box_ai_client)
            
            # Extract text content for validation
            original_content = "\n".join([p.text for p in doc.paragraphs])
            converted_content = "\n".join([p.text for p in converted_doc.paragraphs])
            
            validation_results = validator.validate_conversion(
                original_content, 
                converted_content,
                conga_tags
            )
            
            # Store results in session state
            st.session_state.converted_docs[uploaded_file.name] = {
                'doc': converted_doc,
                'path': exported_path
            }
            
            st.session_state.validation_results[uploaded_file.name] = validation_results
            
        except Exception as e:
            st.error(f"Error converting {uploaded_file.name}: {str(e)}")
        finally:
            # Clean up temporary file
            if os.path.exists(tmp_path):
                os.unlink(tmp_path)
    
    # Complete the progress bar
    progress_bar.progress(1.0)
    status_text.text(f"Completed converting {len(uploaded_files)} files")
    
    # Create a zip file of all converted documents
    if len(uploaded_files) > 1:
        import shutil
        zip_path = os.path.join(tempfile.mkdtemp(), "converted_templates.zip")
        shutil.make_archive(zip_path[:-4], 'zip', output_dir)
        
        with open(zip_path, "rb") as file:
            st.download_button(
                label="Download All Converted Templates (ZIP)",
                data=file,
                file_name="converted_templates.zip",
                mime="application/zip"
            )


def preview_docx(uploaded_file) -> None:
    """
    Preview a DOCX file in Streamlit
    
    Args:
        uploaded_file: Uploaded file object
    """
    try:
        doc = docx.Document(io.BytesIO(uploaded_file.getvalue()))
        
        # Display paragraphs
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                st.write(paragraph.text)
        
        # Display tables
        for i, table in enumerate(doc.tables):
            st.write(f"Table {i+1}:")
            table_data = []
            
            for row in table.rows:
                row_data = [cell.text for cell in row.cells]
                table_data.append(row_data)
            
            if table_data:
                st.table(table_data)
    
    except Exception as e:
        st.error(f"Error previewing document: {str(e)}")


def preview_docx_from_path(file_path: str) -> None:
    """
    Preview a DOCX file from a file path
    
    Args:
        file_path: Path to the DOCX file
    """
    try:
        doc = docx.Document(file_path)
        
        # Display paragraphs
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                st.write(paragraph.text)
        
        # Display tables
        for i, table in enumerate(doc.tables):
            st.write(f"Table {i+1}:")
            table_data = []
            
            for row in table.rows:
                row_data = [cell.text for cell in row.cells]
                table_data.append(row_data)
            
            if table_data:
                st.table(table_data)
    
    except Exception as e:
        st.error(f"Error previewing document: {str(e)}")


if __name__ == "__main__":
    main()
