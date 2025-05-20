"""
Conversion engine for transforming Conga templates to Box DocGen format.
"""
from __future__ import annotations
from typing import Dict, List, Optional, Any, TYPE_CHECKING

from boxsdk import Client
from docx import Document

# Use string type hints to avoid circular imports
if TYPE_CHECKING:
    from app.box_ai_client import BoxAIClient


class ConversionEngine:
    """Engine for converting Conga templates to Box DocGen format."""
    
    def __init__(self, box_ai_client: Optional['BoxAIClient'] = None):
        """Initialize the conversion engine.
        
        Args:
            box_ai_client: Optional BoxAIClient instance for AI-powered conversions
        """
        self.box_ai_client = box_ai_client
    
    def convert_text(self, text: str) -> str:
        """Convert Conga merge fields in text to Box DocGen format.
        
        Args:
            text: Input text with Conga merge fields
            
        Returns:
            str: Text with converted merge fields
        """
        # Simple rule-based conversion for common patterns
        # This can be extended with more sophisticated rules
        converted = text
        
        # Convert simple merge fields
        converted = converted.replace('&=', '{{')
        converted = converted.replace('&!', '{{')
        converted = converted.replace('&+', '{{')
        converted = converted.replace('&', '}}')
        
        # Convert IF conditions
        converted = self._convert_conditions(converted)
        
        # Convert loops/tables
        converted = self._convert_loops(converted)
        
        return converted
    
    def _convert_conditions(self, text: str) -> str:
        """Convert Conga IF conditions to Handlebars format.
        
        Args:
            text: Input text with Conga conditions
            
        Returns:
            str: Text with converted conditions
        """
        # This is a simplified example - a real implementation would need to handle
        # more complex conditions and nested structures
        lines = text.split('\n')
        converted_lines = []
        
        for line in lines:
            if line.strip().startswith('{IF '):
                # Convert {IF condition "true" "false"} to {{#if condition}}true{{else}}false{{/if}}
                parts = line.split('"')
                if len(parts) >= 5:
                    condition = parts[0].replace('{IF', '').strip()
                    true_value = parts[1]
                    false_value = parts[3] if len(parts) > 3 else ''
                    
                    converted = f'{{{{#if {condition}}}}}'
                    converted += true_value
                    if false_value:
                        converted += f'{{{{else}}}}{false_value}'
                    converted += '{{/if}}'
                    
                    converted_lines.append(converted)
                    continue
            
            converted_lines.append(line)
        
        return '\n'.join(converted_lines)
    
    def _convert_loops(self, text: str) -> str:
        """Convert Conga loops to Handlebars each blocks.
        
        Args:
            text: Input text with Conga loops
            
        Returns:
            str: Text with converted loops
        """
        # This is a simplified example - a real implementation would need to handle
        # more complex loop structures and nested loops
        lines = text.split('\n')
        converted_lines = []
        loop_stack = []
        
        for line in lines:
            stripped = line.strip()
            
            # Handle TABLE/LOOP start
            if stripped.startswith('{TABLE') or stripped.startswith('{LOOP'):
                # Extract collection name
                collection = ''
                if 'Group=' in stripped:
                    collection = stripped.split('Group=')[1].split('}')[0].strip()
                elif 'Collection=' in stripped:
                    collection = stripped.split('Collection=')[1].split('}')[0].strip()
                
                if collection:
                    loop_stack.append(collection)
                    converted_lines.append(f'{{{{#each {collection}}}}}')
                    continue
            
            # Handle END
            elif stripped.startswith('{END'):
                if loop_stack:
                    collection = loop_stack.pop()
                    converted_lines.append('{{/each}}')
                    continue
            
            converted_lines.append(line)
        
        return '\n'.join(converted_lines)
    
    def convert_document(self, doc: Document) -> Document:
        """Convert a Word document with Conga templates to Box DocGen format.
        
        Args:
            doc: Input Word document
            
        Returns:
            Document: Converted Word document
        """
        # Create a new document for the output
        new_doc = Document()
        
        # Copy document properties
        new_doc.core_properties = doc.core_properties
        
        # Process each paragraph
        for para in doc.paragraphs:
            converted_text = self.convert_text(para.text)
            new_para = new_doc.add_paragraph(converted_text)
            
            # Copy paragraph formatting
            new_para.style = para.style
            
            # Copy runs and formatting
            for run in para.runs:
                new_run = new_para.add_run(run.text)
                new_run.bold = run.bold
                new_run.italic = run.italic
                new_run.underline = run.underline
                new_run.font.name = run.font.name
                new_run.font.size = run.font.size
                new_run.font.color.rgb = run.font.color.rgb if run.font.color.rgb else None
        
        # Process tables
        for table in doc.tables:
            new_table = new_doc.add_table(rows=len(table.rows), cols=len(table.columns))
            
            for i, row in enumerate(table.rows):
                for j, cell in enumerate(row.cells):
                    converted_text = self.convert_text(cell.text)
                    new_table.cell(i, j).text = converted_text
                    
                    # Copy cell formatting
                    # Note: This is a simplified example - you might want to copy more formatting
                    new_table.cell(i, j).width = cell.width
                    
                    # Copy paragraph formatting in the cell
                    for para in cell.paragraphs:
                        new_para = new_table.cell(i, j).add_paragraph()
                        new_para.style = para.style
                        
                        for run in para.runs:
                            new_run = new_para.add_run(run.text)
                            new_run.bold = run.bold
                            new_run.italic = run.italic
        
        return new_doc
