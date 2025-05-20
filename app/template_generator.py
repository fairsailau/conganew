"""
Module for generating Box DocGen templates from various sources.
"""
from typing import Dict, List, Optional, Any, Union
from pathlib import Path

from docx import Document
from docx.document import Document as DocumentType
from docx.shared import Pt, RGBColor


class DocGenTemplateGenerator:
    """Generator for Box DocGen templates."""
    
    def __init__(self):
        """Initialize the template generator."""
        self.document = Document()
    
    def create_from_ai_output(self, ai_output: str) -> DocumentType:
        """Create a document from AI-generated output.
        
        Args:
            ai_output: String containing the AI-generated content
            
        Returns:
            Document: A Word document with the generated content
        """
        # Create a new document
        doc = Document()
        
        # Add title
        title = doc.add_heading('Box DocGen Template', level=0)
        title.alignment = 1  # Center alignment
        
        # Add a separator
        doc.add_paragraph('_' * 50)
        
        # Add the AI-generated content
        doc.add_paragraph(ai_output)
        
        return doc
    
    def save_to_file(self, doc: DocumentType, file_path: Union[str, Path]) -> None:
        """Save the document to a file.
        
        Args:
            doc: The document to save
            file_path: Path where to save the document
        """
        doc.save(file_path)
    
    def add_section(self, title: str, level: int = 1) -> None:
        """Add a section to the document.
        
        Args:
            title: Section title
            level: Heading level (1-9)
        """
        self.document.add_heading(title, level=min(level, 9))
    
    def add_paragraph(self, text: str, style: Optional[str] = None) -> None:
        """Add a paragraph to the document.
        
        Args:
            text: Paragraph text
            style: Optional style name
        """
        if style:
            self.document.add_paragraph(text, style=style)
        else:
            self.document.add_paragraph(text)
    
    def add_table(self, data: List[List[Any]], 
                 header: bool = True, 
                 style: str = 'Table Grid') -> None:
        """Add a table to the document.
        
        Args:
            data: 2D list of table data
            header: Whether the first row is a header
            style: Table style name
        """
        if not data:
            return
            
        rows = len(data)
        cols = len(data[0]) if rows > 0 else 0
        
        if rows == 0 or cols == 0:
            return
            
        table = self.document.add_table(rows=rows, cols=cols)
        table.style = style
        
        for i, row in enumerate(data):
            for j, cell in enumerate(row):
                table.cell(i, j).text = str(cell)
                
                # Apply header formatting if this is a header row
                if header and i == 0:
                    table.cell(i, j).paragraphs[0].runs[0].bold = True
    
    def add_merge_field(self, field_name: str, description: str = '') -> None:
        """Add a merge field to the document.
        
        Args:
            field_name: Name of the merge field
            description: Optional description of the field
        """
        para = self.document.add_paragraph()
        para.add_run(field_name).bold = True
        if description:
            para.add_run(f": {description}")
    
    def add_conditional_section(self, condition: str, content: str) -> None:
        """Add a conditional section to the document.
        
        Args:
            condition: Condition expression
            content: Content to include if condition is true
        """
        self.document.add_paragraph(f"{{{{#if {condition}}}}}", style='Intense Quote')
        self.document.add_paragraph(content)
        self.document.add_paragraph("{{/if}}", style='Intense Quote')
    
    def add_repeating_section(self, collection_name: str, fields: List[Dict[str, str]]) -> None:
        """Add a repeating section (table) to the document.
        
        Args:
            collection_name: Name of the collection to iterate over
            fields: List of field definitions with 'name' and 'description' keys
        """
        self.document.add_paragraph(f"{{{{#each {collection_name}}}}}", style='Intense Quote')
        
        # Add a table for the repeating section
        table_data = [['Field', 'Value']]
        for field in fields:
            table_data.append([field['name'], f"{{{{this.{field['name']}}}}}"])
        
        self.add_table(table_data, header=True)
        self.document.add_paragraph("{{/each}}", style='Intense Quote')
    
    def add_page_break(self) -> None:
        """Add a page break to the document."""
        self.document.add_page_break()
    
    def get_document(self) -> DocumentType:
        """Get the underlying python-docx Document object.
        
        Returns:
            Document: The python-docx Document object
        """
        return self.document
    
    def clear(self) -> None:
        """Clear the current document content."""
        self.document = Document()
